from io import BytesIO

from docx import Document

from careerpilot.services.parser.detectors import SectionDetector
from careerpilot.services.parser.extractors import DocxExtractor, TxtExtractor
from careerpilot.services.parser.models import ExtractedFact
from careerpilot.services.parser.pipeline import ResumeParserPipeline
from careerpilot.services.parser.validators import FactValidator


def test_extractors_read_txt_and_docx() -> None:
    assert TxtExtractor().extract(b"Skills\nPython") == "Skills\nPython"

    document = Document()
    document.add_heading("Skills")
    document.add_paragraph("Python, SQL")
    content = BytesIO()
    document.save(content)
    assert "Python, SQL" in DocxExtractor().extract(content.getvalue())


def test_section_detection_does_not_depend_on_order() -> None:
    sections = SectionDetector().detect("SKILLS\nPython\nSUMMARY\nBuilder\nEDUCATION\nBS | Test")
    assert sections["skills"] == ["Python"]
    assert sections["summary"] == ["Builder"]
    assert sections["education"] == ["BS | Test"]


def test_pipeline_extracts_normalized_facts() -> None:
    result = ResumeParserPipeline().run(
        "resume.txt",
        "text/plain",
        b"""
SKILLS
Python, SQL
EXPERIENCE
Engineer | Example Corp | 2020 - 2022
EDUCATION
BS Computer Engineering | State University | 2016 - 2020
PROJECTS
Resume parser
""",
    )
    types = [fact.entity_type for fact in result.facts]
    assert types.count("skill") == 2
    assert "experience" in types
    assert "education" in types
    assert "project" in types


def test_validator_reports_duplicates_and_overlaps() -> None:
    facts = [
        ExtractedFact("skill", {"name": "Python"}, 1, "skills:1"),
        ExtractedFact("skill", {"name": "python"}, 1, "skills:2"),
        ExtractedFact(
            "experience",
            {
                "company": "One",
                "title": "Engineer",
                "start_date": "2020-01-01",
                "end_date": "2022-01-01",
            },
            1,
            "experience:1",
        ),
        ExtractedFact(
            "experience",
            {
                "company": "Two",
                "title": "Engineer",
                "start_date": "2021-01-01",
                "end_date": "2023-01-01",
            },
            1,
            "experience:2",
        ),
    ]
    warnings = FactValidator().validate(facts)
    assert any("Duplicate skill" in warning for warning in warnings)
    assert any("Overlapping employment" in warning for warning in warnings)
