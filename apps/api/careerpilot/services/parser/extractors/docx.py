from io import BytesIO

from docx import Document

from careerpilot.services.parser.extractors.base import BaseExtractor, ResumeExtractionError


class DocxExtractor(BaseExtractor):
    mime_types = frozenset(
        {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
    )
    extensions = frozenset({".docx"})

    def extract(self, content: bytes) -> str:
        try:
            document = Document(BytesIO(content))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            table_rows = [
                " | ".join(cell.text for cell in row.cells)
                for table in document.tables
                for row in table.rows
            ]
            return "\n".join([*paragraphs, *table_rows])
        except Exception as error:
            raise ResumeExtractionError("The DOCX file could not be read") from error
