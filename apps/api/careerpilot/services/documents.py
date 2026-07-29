import io
import re
from typing import Any

from docx import Document

from careerpilot.models import CareerProfile, DocumentVersion, Job, JobMatch, ScreeningAnswer
from careerpilot.repositories import DocumentRepository
from careerpilot.services.ai import AIProvider

PROMPT_VERSION = "m6-v1"
SENSITIVE = ("sponsorship", "authorized to work", "salary", "disability", "veteran",
             "race", "gender", "criminal", "clearance", "signature")


class DocumentService:
    def __init__(self, repository: DocumentRepository, provider: AIProvider | None = None) -> None:
        self.repository = repository
        self.provider = provider

    def tailor_resume(
        self, profile: CareerProfile, job: Job, match: JobMatch, template_id=None, use_ai=True
    ) -> DocumentVersion:
        evidence = self._evidence(profile)
        sections = self._base_resume(profile, match)
        if use_ai and self.provider:
            schema = {
                "type": "object",
                "properties": {
                    "summary": {"type": "string"},
                    "experience_bullets": {"type": "array", "items": {"type": "string"}},
                },
                "required": ["summary", "experience_bullets"],
            }
            prompt = (
                "Rewrite only the supplied verified facts for an ATS resume. Do not add facts, "
                "skills, metrics, employers, titles, or dates. Return JSON.\n"
                f"JOB:\n{job.title}\n{job.description}\nFACTS:\n{sections}"
            )
            suggestion = self.provider.generate_json(prompt, schema)
            sections["summary"] = suggestion.get("summary", sections["summary"])
            bullets = suggestion.get("experience_bullets", [])
            if bullets:
                sections["suggested_bullets"] = bullets
        content_text = " ".join(self._flatten(sections)).casefold()
        matched = [item["text"] for item in match.evidence if item.get("type") == "skill"]
        coverage = {
            "matched": [item for item in matched if item.casefold() in content_text],
            "missing": [item for item in matched if item.casefold() not in content_text],
            "ats_safe": True,
        }
        value = DocumentVersion(
            profile_id=profile.id,
            job_id=job.id,
            template_id=template_id,
            document_type="resume",
            version=self.repository.next_version(profile.id, "resume", job.id),
            status="draft",
            title=f"{profile.first_name} {profile.last_name} — {job.title}",
            content=sections,
            evidence=evidence,
            keyword_coverage=coverage,
            model=getattr(self.provider, "model", None) if use_ai else None,
            prompt_version=PROMPT_VERSION,
        )
        return self.repository.save(value)

    def cover_letter(
        self, profile: CareerProfile, job: Job, match: JobMatch, tone: str, use_ai: bool
    ) -> DocumentVersion:
        facts = self._evidence(profile)
        body = (
            f"Dear Hiring Team,\n\nI am interested in the {job.title} opportunity. "
            "My verified experience aligns with "
            f"{', '.join(match.strengths[:3]) or 'the role'}.\n\n"
            "Thank you for your consideration."
        )
        if use_ai and self.provider:
            schema = {
                "type": "object",
                "properties": {"body": {"type": "string"}},
                "required": ["body"],
            }
            body = self.provider.generate_json(
                "Write a concise cover letter using only these verified facts. "
                f"Tone: {tone}. Job: {job.title}. Facts: {facts}", schema
            )["body"]
        return self.repository.save(
            DocumentVersion(
                profile_id=profile.id,
                job_id=job.id,
                document_type="cover_letter",
                version=self.repository.next_version(profile.id, "cover_letter", job.id),
                status="draft",
                title=f"Cover letter — {job.title}",
                content={"body": body, "tone": tone},
                evidence=facts,
                keyword_coverage={},
                model=getattr(self.provider, "model", None) if use_ai else None,
                prompt_version=PROMPT_VERSION,
            )
        )

    def screening(
        self, profile: CareerProfile, job: Job, questions: list[str], use_ai: bool
    ) -> list[ScreeningAnswer]:
        evidence = self._evidence(profile)
        values = []
        for question in questions:
            sensitive = any(term in question.casefold() for term in SENSITIVE)
            answer = None
            confidence = 0.0
            if not sensitive and use_ai and self.provider:
                schema = {
                    "type": "object",
                    "properties": {"answer": {"type": "string"}},
                    "required": ["answer"],
                }
                answer = self.provider.generate_json(
                    f"Answer using only verified facts. Question: {question}. Facts: {evidence}",
                    schema,
                )["answer"]
                confidence = 0.75
            values.append(
                ScreeningAnswer(
                    profile_id=profile.id,
                    job_id=job.id,
                    question=question,
                    normalized_question=re.sub(r"\W+", " ", question.casefold()).strip()[:200],
                    answer=answer,
                    evidence=evidence if answer else [],
                    confidence=confidence,
                    sensitive=sensitive,
                    status="needs_review",
                )
            )
        return self.repository.save_answers(values)

    def export(self, value: DocumentVersion, format_name: str) -> bytes:
        lines = [value.title, *self._flatten(value.content)]
        if format_name == "docx":
            output = io.BytesIO()
            document = Document()
            document.add_heading(value.title, 0)
            for line in lines[1:]:
                document.add_paragraph(line)
            document.save(output)
            return output.getvalue()
        if format_name != "pdf":
            raise ValueError("format must be pdf or docx")
        return self._simple_pdf(lines)

    @staticmethod
    def compare(left: DocumentVersion, right: DocumentVersion) -> dict[str, list[str]]:
        left_lines = set(DocumentService._flatten(left.content))
        right_lines = set(DocumentService._flatten(right.content))
        return {
            "added": sorted(right_lines - left_lines),
            "removed": sorted(left_lines - right_lines),
            "unchanged": sorted(left_lines & right_lines),
        }

    @staticmethod
    def _base_resume(profile: CareerProfile, match: JobMatch) -> dict[str, Any]:
        matched = {item["text"].casefold() for item in match.evidence if item.get("text")}
        skills = sorted(profile.skills, key=lambda item: item.name.casefold() not in matched)
        return {
            "contact": {
                "name": f"{profile.first_name} {profile.last_name}",
                "email": profile.email,
                "phone": profile.phone,
                "location": ", ".join(
                    filter(None, [profile.city, profile.region, profile.country])
                ),
            },
            "summary": f"{profile.experiences[0].title} with verified experience."
            if profile.experiences else "Professional profile",
            "skills": [item.name for item in skills],
            "experience": [
                {
                    "company": item.company,
                    "title": item.title,
                    "dates": f"{item.start_date} — {item.end_date or 'Present'}",
                    "description": item.description,
                }
                for item in profile.experiences
            ],
            "projects": [
                {"name": item.name, "description": item.description}
                for item in profile.projects
            ],
            "education": [
                {
                    "institution": item.institution,
                    "degree": item.degree,
                    "field": item.field_of_study,
                }
                for item in profile.education
            ],
        }

    @staticmethod
    def _evidence(profile: CareerProfile) -> list[dict[str, Any]]:
        values = []
        for collection, kind, text in (
            (profile.skills, "skill", lambda x: x.name),
            (
                profile.experiences,
                "experience",
                lambda x: f"{x.title} at {x.company}: {x.description or ''}",
            ),
            (profile.projects, "project", lambda x: f"{x.name}: {x.description or ''}"),
            (profile.achievements, "achievement", lambda x: f"{x.title}: {x.description}"),
        ):
            values.extend(
                {"type": kind, "id": str(item.id), "text": text(item)}
                for item in collection
            )
        return values

    @staticmethod
    def _flatten(value: Any) -> list[str]:
        if isinstance(value, dict):
            return [line for child in value.values() for line in DocumentService._flatten(child)]
        if isinstance(value, list):
            return [line for child in value for line in DocumentService._flatten(child)]
        return [str(value)] if value not in (None, "") else []

    @staticmethod
    def _simple_pdf(lines: list[str]) -> bytes:
        escaped = [
            line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")[:100]
            for line in lines
        ]
        stream = "BT /F1 10 Tf 50 760 Td 14 TL " + " ".join(
            f"({line}) Tj T*" for line in escaped[:48]
        ) + " ET"
        objects = [
            "<< /Type /Catalog /Pages 2 0 R >>",
            "<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
            "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            "/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
            f"<< /Length {len(stream.encode())} >>\nstream\n{stream}\nendstream",
            "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        ]
        result = bytearray(b"%PDF-1.4\n")
        offsets = [0]
        for index, obj in enumerate(objects, 1):
            offsets.append(len(result))
            result.extend(f"{index} 0 obj\n{obj}\nendobj\n".encode())
        xref = len(result)
        result.extend(f"xref\n0 {len(objects)+1}\n0000000000 65535 f \n".encode())
        for offset in offsets[1:]:
            result.extend(f"{offset:010d} 00000 n \n".encode())
        result.extend(
            f"trailer << /Size {len(objects)+1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF".encode()
        )
        return bytes(result)
